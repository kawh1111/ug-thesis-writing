#!/usr/bin/env python3
"""
批量修正 Word (.docx) 文件的格式问题。

用法:
    python apply_format_fixes.py <docx_path> [--config format.json] [--output fixed.docx]
    python apply_format_fixes.py <docx_path> --preview  # 预览模式，不实际修改

功能:
    - 修正正文字号、字体、行距、首行缩进、对齐方式
    - 修正标题格式
    - 修正图表名格式

安全措施:
    - 自动创建备份
    - Zotero 字段计数验证
    - 输出为新文件，不覆盖原文件
"""

import argparse
import json
import shutil
import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("错误: 需要安装 python-docx。运行: pip install python-docx", file=sys.stderr)
    sys.exit(1)


# 默认格式配置
DEFAULT_CONFIG = {
    "body": {
        "font_cn": "宋体",
        "font_en": "Times New Roman",
        "size_pt": 12,
        "line_spacing": 1.5,
        "first_indent_chars": 2,
        "alignment": "both",
    },
    "heading1": {
        "font_cn": "黑体",
        "size_pt": 16,
        "bold": True,
        "alignment": "center",
    },
    "heading2": {
        "font_cn": "黑体",
        "size_pt": 14,
        "bold": True,
        "alignment": "left",
    },
    "heading3": {
        "font_cn": "宋体",
        "size_pt": 12,
        "bold": False,
        "alignment": "left",
    },
    "caption": {
        "font_cn": "宋体",
        "size_pt": 10.5,
        "alignment": "center",
    },
}


ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "both": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def count_zotero_fields(doc: Document) -> int:
    """计数文档中的 Zotero 字段数量。"""
    count = 0
    for para in doc.paragraphs:
        if "ADDIN ZOTERO" in para.text or "ZOTERO_ITEM" in para.text:
            count += 1
    # 也检查页眉页脚
    for section in doc.sections:
        for header_footer in [section.header, section.footer]:
            if header_footer:
                for para in header_footer.paragraphs:
                    if "ADDIN ZOTERO" in para.text or "ZOTERO_ITEM" in para.text:
                        count += 1
    return count


def identify_paragraph_type(para) -> str:
    """识别段落类型。"""
    style_name = para.style.name if para.style else ""

    if "Heading 1" in style_name or "标题 1" in style_name:
        return "heading1"
    elif "Heading 2" in style_name or "标题 2" in style_name:
        return "heading2"
    elif "Heading 3" in style_name or "标题 3" in style_name:
        return "heading3"
    elif "TOC" in style_name or "目录" in style_name:
        return "toc"
    elif "Caption" in style_name or "题注" in style_name:
        return "caption"
    elif "Header" in style_name or "Footer" in style_name:
        return "header_footer"
    else:
        return "body"


def has_zotero_field(para) -> bool:
    """检查段落是否包含 Zotero 字段。"""
    return "ADDIN ZOTERO" in para.text or "ZOTERO_ITEM" in para.text


def has_images(para) -> bool:
    """检查段落是否包含图片。"""
    for run in para.runs:
        if run.element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"):
            return True
        if run.element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict"):
            return True
    return False


def has_formulas(para) -> bool:
    """检查段落是否包含公式（OLE 对象）。"""
    for run in para.runs:
        if run.element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}object"):
            return True
    return False


def apply_format_to_paragraph(para, config: dict, fix_types: list[str] | None = None) -> list[str]:
    """修正单个段落的格式。

    Args:
        para: 段落对象
        config: 格式配置
        fix_types: 要修正的类型列表（None 表示全部）

    Returns:
        修正操作列表
    """
    changes = []

    # 跳过含特殊内容的段落
    if has_zotero_field(para):
        changes.append("跳过: 含 Zotero 字段")
        return changes
    if has_images(para):
        changes.append("跳过: 含图片")
        return changes
    if has_formulas(para):
        changes.append("跳过: 含公式")
        return changes

    fix_all = fix_types is None

    # 修正字号
    if fix_all or "size" in fix_types:
        expected_size = config.get("size_pt")
        if expected_size:
            for run in para.runs:
                if run.text.strip():
                    run.font.size = Pt(expected_size)
            changes.append(f"字号 → {expected_size}pt")

    # 修正字体
    if fix_all or "font" in fix_types:
        font_cn = config.get("font_cn")
        font_en = config.get("font_en")
        for run in para.runs:
            if run.text.strip():
                if font_en:
                    run.font.name = font_en
                if font_cn:
                    # 设置东亚字体
                    rPr = run.font.element.get_or_add_rPr()
                    rFonts = rPr.get_or_add_rFonts()
                    rFonts.set(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia",
                        font_cn
                    )
        if font_cn or font_en:
            changes.append(f"字体 → {font_cn or ''}/{font_en or ''}")

    # 修正行距
    if fix_all or "spacing" in fix_types:
        expected_spacing = config.get("line_spacing")
        if expected_spacing:
            para.paragraph_format.line_spacing = expected_spacing
            changes.append(f"行距 → {expected_spacing}倍")

    # 修正首行缩进
    if fix_all or "indent" in fix_types:
        expected_indent = config.get("first_indent_chars")
        if expected_indent:
            para.paragraph_format.first_line_indent = Cm(expected_indent * 0.37)
            changes.append(f"首行缩进 → {expected_indent}字符")

    # 修正对齐方式
    if fix_all or "alignment" in fix_types:
        expected_align = config.get("alignment")
        if expected_align and expected_align in ALIGNMENT_MAP:
            para.paragraph_format.alignment = ALIGNMENT_MAP[expected_align]
            changes.append(f"对齐 → {expected_align}")

    # 修正加粗
    if "bold" in config:
        for run in para.runs:
            if run.text.strip():
                run.font.bold = config["bold"]

    return changes


def apply_fixes(
    docx_path: str,
    config: dict | None = None,
    output_path: str | None = None,
    preview: bool = False,
    fix_types: list[str] | None = None,
) -> dict:
    """批量修正 DOCX 文件格式。

    Args:
        docx_path: DOCX 文件路径
        config: 格式配置
        output_path: 输出文件路径
        preview: 预览模式（不实际修改）
        fix_types: 要修正的类型列表

    Returns:
        修正结果统计
    """
    if config is None:
        config = DEFAULT_CONFIG

    # 创建备份
    if not preview:
        backup_path = f"{docx_path}.backup_{datetime.now():%Y%m%d_%H%M%S}"
        shutil.copy2(docx_path, backup_path)
        print(f"备份已创建: {backup_path}")

    doc = Document(docx_path)

    # 编辑前 Zotero 字段计数
    zotero_before = count_zotero_fields(doc)

    stats = {
        "total_paragraphs": len(doc.paragraphs),
        "modified_paragraphs": 0,
        "skipped_paragraphs": 0,
        "changes": [],
    }

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        para_type = identify_paragraph_type(para)

        # 跳过不需要修改的类型
        if para_type in ("toc", "header_footer"):
            stats["skipped_paragraphs"] += 1
            continue

        # 获取对应的格式配置
        type_config = config.get(para_type, config.get("body", {}))
        if not type_config:
            stats["skipped_paragraphs"] += 1
            continue

        if preview:
            # 预览模式：只记录需要修改的内容
            stats["changes"].append({
                "paragraph": i,
                "type": para_type,
                "text_preview": text[:50],
                "config": type_config,
            })
        else:
            # 实际修改
            changes = apply_format_to_paragraph(para, type_config, fix_types)
            if changes:
                stats["modified_paragraphs"] += 1
                stats["changes"].append({
                    "paragraph": i,
                    "type": para_type,
                    "text_preview": text[:50],
                    "operations": changes,
                })

    # 编辑后 Zotero 字段计数
    if not preview:
        zotero_after = count_zotero_fields(doc)
        if zotero_before != zotero_after:
            print(f"警告: Zotero 字段数量变化！编辑前: {zotero_before}, 编辑后: {zotero_after}")
            print("建议检查引用是否完整。")
        else:
            print(f"Zotero 字段验证通过: {zotero_before} 个")

        # 保存
        if output_path is None:
            stem = Path(docx_path).stem
            suffix = Path(docx_path).suffix
            output_path = str(Path(docx_path).parent / f"{stem}_格式修正{suffix}")

        doc.save(output_path)
        print(f"修正后文件已保存: {output_path}")

    return stats


def main():
    parser = argparse.ArgumentParser(
        description="批量修正 Word (.docx) 文件的格式问题"
    )
    parser.add_argument("docx_path", help="DOCX 文件路径")
    parser.add_argument("--config", "-c", help="格式配置文件 (JSON)")
    parser.add_argument("--output", "-o", help="输出文件路径")
    parser.add_argument("--preview", action="store_true", help="预览模式，不实际修改")
    parser.add_argument("--fix-types", nargs="+",
                       choices=["size", "font", "spacing", "indent", "alignment"],
                       help="只修正指定的类型")

    args = parser.parse_args()

    docx_path = Path(args.docx_path)
    if not docx_path.exists():
        print(f"错误: 文件不存在: {docx_path}", file=sys.stderr)
        sys.exit(1)

    # 加载配置
    config = None
    if args.config:
        config_path = Path(args.config)
        if not config_path.exists():
            print(f"错误: 配置文件不存在: {config_path}", file=sys.stderr)
            sys.exit(1)
        config = json.loads(config_path.read_text(encoding="utf-8"))

    # 执行修正
    stats = apply_fixes(
        str(docx_path),
        config=config,
        output_path=args.output,
        preview=args.preview,
        fix_types=args.fix_types,
    )

    # 输出统计
    print(f"\n--- 统计 ---")
    print(f"总段落数: {stats['total_paragraphs']}")
    print(f"已修改: {stats['modified_paragraphs']}")
    print(f"已跳过: {stats['skipped_paragraphs']}")

    if args.preview and stats["changes"]:
        print(f"\n需要修改的段落:")
        for change in stats["changes"][:20]:  # 只显示前20个
            print(f"  P{change['paragraph']} [{change['type']}] {change['text_preview']}...")
        if len(stats["changes"]) > 20:
            print(f"  ... 还有 {len(stats['changes']) - 20} 个段落")


if __name__ == "__main__":
    main()
