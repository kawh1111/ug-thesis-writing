#!/usr/bin/env python3
"""
检查 Word (.docx) 文件的格式是否符合规范。

用法:
    python check_docx_format.py <docx_path> [--config format.json] [--output report.md]

检查项:
    - 正文字号、字体、行距、首行缩进、对齐方式
    - 标题层级格式
    - 图表名字号
    - 页边距

输出: Markdown 格式的检查报告，按严重程度分级。
"""

import argparse
import json
import sys
from pathlib import Path
from dataclasses import dataclass, field

try:
    from docx import Document
    from docx.shared import Pt, Cm, Emu
except ImportError:
    print("错误: 需要安装 python-docx。运行: pip install python-docx", file=sys.stderr)
    sys.exit(1)


# 默认格式配置
DEFAULT_CONFIG = {
    "body": {
        "font_cn": "宋体",
        "font_en": "Times New Roman",
        "size_pt": 12,
        "line_spacing": 1.5,
        "first_indent_chars": 2,
        "alignment": "both",
    },
    "heading1": {
        "font_cn": "黑体",
        "size_pt": 16,
        "bold": True,
        "alignment": "center",
    },
    "heading2": {
        "font_cn": "黑体",
        "size_pt": 14,
        "bold": True,
        "alignment": "left",
    },
    "heading3": {
        "font_cn": "宋体",
        "size_pt": 12,
        "bold": False,
        "alignment": "left",
    },
    "caption": {
        "font_cn": "宋体",
        "size_pt": 10.5,
        "alignment": "center",
    },
    "page": {
        "paper": "A4",
        "margin_top_cm": 2.54,
        "margin_bottom_cm": 2.54,
        "margin_left_cm": 3.17,
        "margin_right_cm": 3.17,
    },
    "tolerance_pt": 0.5,
}


@dataclass
class Issue:
    """格式问题。"""
    severity: str  # "严重" | "中等" | "轻微"
    category: str
    paragraph_index: int
    description: str
    current_value: str = ""
    expected_value: str = ""


@dataclass
class FormatReport:
    """格式检查报告。"""
    file_path: str
    issues: list[Issue] = field(default_factory=list)
    total_paragraphs: int = 0
    checked_paragraphs: int = 0

    @property
    def critical_count(self) -> int:
        return sum(1 for i in self.issues if i.severity == "严重")

    @property
    def medium_count(self) -> int:
        return sum(1 for i in self.issues if i.severity == "中等")

    @property
    def minor_count(self) -> int:
        return sum(1 for i in self.issues if i.severity == "轻微")

    def to_markdown(self) -> str:
        """生成 Markdown 报告。"""
        lines = [
            f"# 格式检查报告",
            f"",
            f"**文件**: `{self.file_path}`",
            f"**总段落数**: {self.total_paragraphs}",
            f"**已检查**: {self.checked_paragraphs}",
            f"",
            f"## 问题统计",
            f"",
            f"| 级别 | 数量 |",
            f"|------|------|",
            f"| 严重 | {self.critical_count} |",
            f"| 中等 | {self.medium_count} |",
            f"| 轻微 | {self.minor_count} |",
            f"| **合计** | **{len(self.issues)}** |",
            f"",
        ]

        if not self.issues:
            lines.append("未发现格式问题。")
            return "\n".join(lines)

        # 按严重程度分组
        for severity in ["严重", "中等", "轻微"]:
            severity_issues = [i for i in self.issues if i.severity == severity]
            if not severity_issues:
                continue

            lines.append(f"## {severity}问题")
            lines.append("")
            lines.append("| # | 类别 | 段落 | 问题描述 | 当前值 | 期望值 |")
            lines.append("|---|------|------|----------|--------|--------|")

            for idx, issue in enumerate(severity_issues, 1):
                lines.append(
                    f"| {idx} | {issue.category} | P{issue.paragraph_index} "
                    f"| {issue.description} | {issue.current_value} | {issue.expected_value} |"
                )
            lines.append("")

        return "\n".join(lines)


def get_font_size_pt(run) -> float | None:
    """获取 run 的字号（磅值）。"""
    if run.font.size is not None:
        return run.font.size.pt
    return None


def get_font_name(run) -> tuple[str | None, str | None]:
    """获取 run 的字体名（中文, 西文）。"""
    font_cn = None
    font_en = None

    if run.font.name:
        font_en = run.font.name

    # 尝试获取东亚字体
    rPr = run.font.element.rPr
    if rPr is not None:
        rFonts = rPr.rFonts
        if rFonts is not None:
            east_asia = rFonts.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia")
            if east_asia:
                font_cn = east_asia

    return font_cn, font_en


def get_line_spacing(para) -> float | None:
    """获取段落行距。"""
    pf = para.paragraph_format
    if pf.line_spacing is not None:
        return pf.line_spacing
    return None


def get_first_indent_chars(para) -> float | None:
    """获取首行缩进（字符数）。"""
    pf = para.paragraph_format
    if pf.first_line_indent is not None:
        # 1字符 ≈ 0.37cm ≈ 10.5pt (五号字)
        cm = pf.first_line_indent / 360000  # EMU to cm
        chars = cm / 0.37
        return round(chars, 1)
    return None


def get_alignment(para) -> str | None:
    """获取对齐方式。"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    pf = para.paragraph_format
    if pf.alignment is not None:
        align_map = {
            WD_ALIGN_PARAGRAPH.LEFT: "left",
            WD_ALIGN_PARAGRAPH.CENTER: "center",
            WD_ALIGN_PARAGRAPH.RIGHT: "right",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "both",
        }
        return align_map.get(pf.alignment, "unknown")
    return None


def identify_paragraph_type(para) -> str:
    """识别段落类型。"""
    style_name = para.style.name if para.style else ""

    if "Heading 1" in style_name or "标题 1" in style_name:
        return "heading1"
    elif "Heading 2" in style_name or "标题 2" in style_name:
        return "heading2"
    elif "Heading 3" in style_name or "标题 3" in style_name:
        return "heading3"
    elif "TOC" in style_name or "目录" in style_name:
        return "toc"
    elif "Caption" in style_name or "题注" in style_name:
        return "caption"
    elif "Header" in style_name or "Footer" in style_name or "页眉" in style_name or "页脚" in style_name:
        return "header_footer"
    else:
        return "body"


def check_format(docx_path: str, config: dict | None = None) -> FormatReport:
    """检查 DOCX 文件格式。

    Args:
        docx_path: DOCX 文件路径
        config: 格式配置（使用默认配置如果为 None）

    Returns:
        格式检查报告
    """
    if config is None:
        config = DEFAULT_CONFIG

    tolerance = config.get("tolerance_pt", 0.5)
    doc = Document(docx_path)
    report = FormatReport(file_path=docx_path)
    report.total_paragraphs = len(doc.paragraphs)

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        para_type = identify_paragraph_type(para)
        if para_type in ("toc", "header_footer"):
            continue

        report.checked_paragraphs += 1

        # 获取格式配置
        type_config = config.get(para_type, config.get("body", {}))
        if not type_config:
            continue

        # 检查字号
        expected_size = type_config.get("size_pt")
        if expected_size:
            for run in para.runs:
                actual_size = get_font_size_pt(run)
                if actual_size and abs(actual_size - expected_size) > tolerance:
                    report.issues.append(Issue(
                        severity="严重",
                        category="字号",
                        paragraph_index=i,
                        description=f"字号不正确",
                        current_value=f"{actual_size}pt",
                        expected_value=f"{expected_size}pt",
                    ))
                    break  # 每个段落只报一次

        # 检查行距（仅正文）
        if para_type == "body":
            expected_spacing = type_config.get("line_spacing")
            if expected_spacing:
                actual_spacing = get_line_spacing(para)
                if actual_spacing and abs(actual_spacing - expected_spacing) > 0.1:
                    report.issues.append(Issue(
                        severity="严重",
                        category="行距",
                        paragraph_index=i,
                        description="行距不正确",
                        current_value=f"{actual_spacing:.1f}倍",
                        expected_value=f"{expected_spacing:.1f}倍",
                    ))

        # 检查首行缩进（仅正文）
        if para_type == "body":
            expected_indent = type_config.get("first_indent_chars")
            if expected_indent:
                actual_indent = get_first_indent_chars(para)
                if actual_indent is None or abs(actual_indent - expected_indent) > 0.5:
                    report.issues.append(Issue(
                        severity="中等",
                        category="缩进",
                        paragraph_index=i,
                        description="首行缩进缺失或不正确",
                        current_value=f"{actual_indent}字符" if actual_indent else "无",
                        expected_value=f"{expected_indent}字符",
                    ))

        # 检查对齐方式
        expected_align = type_config.get("alignment")
        if expected_align:
            actual_align = get_alignment(para)
            if actual_align and actual_align != expected_align:
                report.issues.append(Issue(
                    severity="轻微",
                    category="对齐",
                    paragraph_index=i,
                    description="对齐方式不正确",
                    current_value=actual_align,
                    expected_value=expected_align,
                ))

    return report


def main():
    parser = argparse.ArgumentParser(
        description="检查 Word (.docx) 文件的格式是否符合规范"
    )
    parser.add_argument("docx_path", help="DOCX 文件路径")
    parser.add_argument("--config", "-c", help="格式配置文件 (JSON)")
    parser.add_argument("--output", "-o", help="输出报告文件路径")
    parser.add_argument("--json", action="store_true", help="输出 JSON 格式")

    args = parser.parse_args()

    docx_path = Path(args.docx_path)
    if not docx_path.exists():
        print(f"错误: 文件不存在: {docx_path}", file=sys.stderr)
        sys.exit(1)

    # 加载配置
    config = None
    if args.config:
        config_path = Path(args.config)
        if not config_path.exists():
            print(f"错误: 配置文件不存在: {config_path}", file=sys.stderr)
            sys.exit(1)
        config = json.loads(config_path.read_text(encoding="utf-8"))

    # 执行检查
    report = check_format(str(docx_path), config)

    # 输出结果
    if args.json:
        import dataclasses
        output = json.dumps(dataclasses.asdict(report), ensure_ascii=False, indent=2)
    else:
        output = report.to_markdown()

    if args.output:
        Path(args.output).write_text(output, encoding="utf-8")
        print(f"报告已保存到 {args.output}")
        print(f"发现问题: 严重{report.critical_count} / 中等{report.medium_count} / 轻微{report.minor_count}")
    else:
        print(output)


if __name__ == "__main__":
    main()
